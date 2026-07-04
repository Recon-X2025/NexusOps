#!/usr/bin/env python3
"""Generate the non-technical, circulatable DOCX competitive gap summary."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/Users/kathikiyer/Documents/NexusOps/docs-word/CoheronConnect_Competitive_Gap_Summary_2026-06-30.docx"

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6E, 0x00)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, color=None, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def rating_run(cell, value):
    p = cell.paragraphs[0]
    r = p.add_run(value)
    r.bold = True
    if value == "GA":
        r.font.color.rgb = GREEN
    elif value == "PARTIAL":
        r.font.color.rgb = AMBER
    else:
        r.font.color.rgb = RED


# ---- Title block ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("CoheronConnect")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
r = st.add_run("Competitive Gap Summary — How We Stack Up Against the Market")
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
r = meta.add_run("Prepared 30 June 2026  •  Based on a full review of the live product against the 2026 capabilities of category leaders  •  Two views: Enterprise and Startups/SMBs (up to 500 employees)")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---- The one-line story ----
heading("The short version", 15)
para("CoheronConnect's strength is breadth on a single platform — it does the work of six or more separate systems (IT service desk, finance, HR & payroll, procurement, CRM, and compliance) for one organisation, with unusually deep support for Indian statutory requirements. Its trade-off is depth: in any single area, the dedicated market leader goes further, and the 2026 expectation of built-in AI assistance is not yet met across the board.")
para("Bottom line: We are competitive today for startups and SMBs up to 500 employees — especially those operating in India — as a 'replace many tools with one' play. For large enterprises, we are suited to a single department or region rather than replacing their core systems, until we close gaps in security administration, reporting depth, scale assurances, and AI.")

# ---- What's genuinely strong ----
heading("What is genuinely strong", 15)
bullet(" Real double-entry accounting with enforced balancing, full procure-to-pay with three-way invoice matching.", "Finance controls:")
bullet(" Deep, India-specific payroll and tax calculations (both tax regimes, surcharge relief, EPF/ESI/PT), plus GST e-invoicing, TDS, compliant SMS, and Aadhaar e-signature.", "India compliance:")
bullet(" Enterprise sign-in (SAML SSO), role-based access, encrypted credentials, and an audit trail.", "Security basics:")
bullet(" One consistent, multi-tenant platform spanning all the modules — the 'one system' story is real in the product, not just on a slide.", "Breadth:")

# ---- The four weaknesses ----
heading("The four things holding us back", 15)
bullet(" Every leader now ships built-in AI (auto-triage, drafting, forecasting, even autonomous actions). We have an AI assistant but not yet AI woven into each module.", "AI expectations:")
bullet(" A few modules (application monitoring, DevOps, on-call paging) currently store records rather than perform the live function buyers expect. These should be re-described or strengthened.", "Some modules look deeper than they are:")
bullet(" We have fixed dashboards but no self-serve report builder or forecasting — a basic expectation, especially at enterprise level.", "Reporting depth:")
bullet(" Missing automated user provisioning (SCIM), self-service two-factor authentication, and finer-grained permissions that large IT teams require.", "Enterprise administration:")

# ---- How to read the table ----
heading("Module-by-module scorecard", 15)
para("Each module is rated for two buyer types. GA = competitive / good enough to win. PARTIAL = usable but with noticeable gaps. GAP = materially behind for that buyer.", italic=True, color=GREY, size=9.5)

rows = [
    ("ITSM / IT Service Desk", "PARTIAL", "GA"),
    ("Asset & Configuration (CMDB)", "GAP", "PARTIAL"),
    ("Application Monitoring / On-call", "GAP", "GAP"),
    ("DevOps / CI-CD", "GAP", "GAP"),
    ("Core HR", "PARTIAL", "GA*"),
    ("Payroll — India", "PARTIAL", "GA"),
    ("Payroll — Global", "GAP", "GAP"),
    ("Accounting / Finance", "GAP", "GA"),
    ("Procurement / Vendors", "PARTIAL", "GA"),
    ("CRM / Sales", "GAP", "PARTIAL"),
    ("Governance, Risk & Compliance", "GAP", "PARTIAL"),
    ("India Compliance / Secretarial / Legal / e-Sign", "PARTIAL", "GA"),
    ("Platform (sign-in, roles, integrations)", "PARTIAL", "GA"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = table.rows[0].cells
for i, htext in enumerate(["Module", "Vs. Enterprise leaders", "Vs. SMB ≤500 needs"]):
    p = hdr[i].paragraphs[0]
    rr = p.add_run(htext)
    rr.bold = True
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for name, ent, smb in rows:
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run(name).font.size = Pt(10)
    rating_run(cells[1], ent)
    rating_run(cells[2], smb)

para("*Good for India-centric SMBs; partial for global teams.", italic=True, color=GREY, size=9)

# ---- Per-segment summary ----
heading("What this means for each buyer", 15)

heading("Startups & SMBs (up to 500 employees)", 12.5, color=ACCENT, space_before=8, space_after=4)
para("We are competitive now. The pitch is consolidation: one platform instead of six tools, at lower total cost and with India compliance built in. To stay ahead we should add per-module AI assistance, a simple report builder, bank reconciliation, and employee self-service — and re-describe the monitoring/DevOps/on-call modules to set accurate expectations.")

heading("Large Enterprises", 12.5, color=ACCENT, space_before=8, space_after=4)
para("We are best positioned as a system for a single department or region (for example, India operations, compliance, or employee service management) rather than replacing a company's core IT, finance, or HR platform. To unlock those core seats we need enterprise sign-in administration (automated provisioning, self-service two-factor), deeper reporting and multi-entity finance, published scale and reliability assurances, and AI parity.")

# ---- Priorities ----
heading("Recommended priorities", 15)
heading("To win more SMB deals", 12.5, color=ACCENT, space_before=8, space_after=4)
for x in [
    "Add AI assistance inside each module (auto-triage tickets, capture invoices, flag payroll anomalies).",
    "Ship a lightweight report builder with saved views and scheduled exports.",
    "Add bank reconciliation (finance) and employee self-service (HR/payroll).",
    "Lead India go-to-market with the compliance + secretarial + e-signature bundle — our clearest advantage.",
]:
    bullet(" " + x)

heading("To unlock enterprise departments", 12.5, color=ACCENT, space_before=8, space_after=4)
for x in [
    "Automated user provisioning (SCIM) and self-service two-factor authentication.",
    "A custom-role builder and finer-grained permissions.",
    "Multi-entity, multi-currency finance consolidation.",
    "Publish scale, reliability and disaster-recovery assurances; pursue SOC 2 / ISO 27001 certification.",
]:
    bullet(" " + x)

doc.add_paragraph()
foot = doc.add_paragraph()
r = foot.add_run("This is a business summary. A detailed, code-referenced version is maintained by Engineering (COMPETITIVE_GAP_ANALYSIS_2026-06-30.md).")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

doc.save(OUT)
print("Wrote", OUT)
