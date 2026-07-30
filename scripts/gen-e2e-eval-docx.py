#!/usr/bin/env python3
"""
gen-e2e-eval-docx.py — Render docs/E2E_EVALUATION_FINDINGS.md into a formatted
Word document with embedded screenshots.

Reads:  docs/E2E_EVALUATION_FINDINGS.md
Images: docs/e2e-screenshots/<name>.png  (referenced by `**Screenshot:** \`name.png\``)
Writes: docs-word/E2E_Evaluation_Findings.docx

Rendering rules:
  # / ## / ### / ####   → Word headings (levels 0..4)
  **Finding:** id lines  → severity-coloured finding heading
  - bullets              → bullet list, **bold** + `code` runs styled inline
  > blockquote           → indented italic note
  ---                    → (skipped; sections are heading-delimited)
  `**Screenshot:** name` → embed docs/e2e-screenshots/name below the bullet
An executive summary (severity tally + module index) is generated up front.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "docs" / "E2E_EVALUATION_FINDINGS.md"
IMG_DIR = ROOT / "docs" / "e2e-screenshots"
OUT_DIR = ROOT / "docs-word"
OUT = OUT_DIR / "E2E_Evaluation_Findings.docx"

SEV_COLORS = {
    "BLOCKER": RGBColor(0xC0, 0x00, 0x00),
    "HIGH":    RGBColor(0xD8, 0x4B, 0x00),
    "MEDIUM":  RGBColor(0xB8, 0x8A, 0x00),
    "LOW":     RGBColor(0x3C, 0x6E, 0x00),
    "INFO":    RGBColor(0x2E, 0x5A, 0x88),
    "RESOLVED":RGBColor(0x3C, 0x6E, 0x00),
}

SCREENSHOT_RE = re.compile(r"\*\*Screenshot:\*\*\s*`([^`]+)`")
SEV_LINE_RE = re.compile(r"\*\*Severity:\*\*\s*([A-Z/ ]+)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text):
    """Add a run sequence, honouring **bold** and `code` spans."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x8B, 0x00, 0x50)
        else:
            paragraph.add_run(part)


def severity_of_block(lines, idx):
    """Look ahead a few lines from a finding heading for its Severity."""
    for j in range(idx, min(idx + 4, len(lines))):
        m = SEV_LINE_RE.search(lines[j])
        if m:
            sev = m.group(1).strip().split("/")[0].split()[0]
            return sev
    return None


def collect_stats(lines):
    """Count findings by severity and by module for the exec summary."""
    tally = {k: 0 for k in ["BLOCKER", "HIGH", "MEDIUM", "LOW", "INFO"]}
    findings = []  # (id, severity, title)
    cur_module = "Cross-cutting"
    for i, ln in enumerate(lines):
        if ln.startswith("### MODULE") or ln.startswith("## MODULE"):
            cur_module = ln.lstrip("#").strip()
        m = re.match(r"^####\s+(F-[A-Z0-9]+)\s+—\s+(.*)$", ln)
        if m:
            fid, title = m.group(1), m.group(2)
            sev = severity_of_block(lines, i) or "INFO"
            base = sev if sev in tally else "INFO"
            tally[base] += 1
            findings.append((fid, base, title, cur_module))
    return tally, findings


def main():
    if not MD.exists():
        print(f"missing {MD}", file=sys.stderr)
        sys.exit(1)
    OUT_DIR.mkdir(exist_ok=True)
    lines = MD.read_text(encoding="utf-8").splitlines()

    doc = Document()
    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    tally, findings = collect_stats(lines)

    # ── Cover ────────────────────────────────────────────────────────────────
    t = doc.add_heading("CoheronConnect (NexusOps)", level=0)
    doc.add_heading("Deep End-to-End Evaluation — Findings Report", level=1)
    p = doc.add_paragraph()
    p.add_run("Per-module vertical drill · positive + negative · live-driven with screenshots").italic = True

    # ── Executive summary ────────────────────────────────────────────────────
    doc.add_heading("Executive summary", level=1)
    p = doc.add_paragraph()
    total = sum(tally.values())
    p.add_run(f"{total} module-level findings recorded across 8 modules "
              f"(plus 5 cross-cutting). Severity distribution:")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Count"
    for sev in ["BLOCKER", "HIGH", "MEDIUM", "LOW", "INFO"]:
        row = tbl.add_row().cells
        rp = row[0].paragraphs[0]
        run = rp.add_run(sev)
        run.bold = True
        if sev in SEV_COLORS:
            run.font.color.rgb = SEV_COLORS[sev]
        row[1].text = str(tally[sev])

    doc.add_heading("Findings index", level=2)
    for fid, sev, title, module in findings:
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(f"{fid} ")
        r.bold = True
        r.font.color.rgb = SEV_COLORS.get(sev, RGBColor(0, 0, 0))
        para.add_run(f"[{sev}] ")
        para.add_run(title)
    doc.add_page_break()

    # ── Body: stream the markdown ────────────────────────────────────────────
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip("\n")
        stripped = ln.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Screenshot embed
        sm = SCREENSHOT_RE.search(ln)
        if sm and ln.lstrip().startswith("- **Screenshot:**"):
            name = sm.group(1)
            img = IMG_DIR / name
            cap = doc.add_paragraph()
            cr = cap.add_run(f"Screenshot: {name}")
            cr.bold = True
            cr.font.size = Pt(9)
            if img.exists():
                try:
                    doc.add_picture(str(img), width=Inches(6.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:  # noqa
                    doc.add_paragraph(f"[could not embed {name}: {e}]")
            else:
                doc.add_paragraph(f"[missing image: {name}]")
            i += 1
            continue

        # Headings
        if ln.startswith("#### "):
            text = ln[5:].strip()
            m = re.match(r"^(F-[A-Z0-9]+)\s+—\s+(.*)$", text)
            h = doc.add_heading(level=4)
            if m:
                fid, title = m.group(1), m.group(2)
                sev = severity_of_block(lines, i) or "INFO"
                r = h.add_run(f"{fid} — {title}")
                r.font.color.rgb = SEV_COLORS.get(sev, RGBColor(0, 0, 0))
            else:
                h.add_run(text)
            i += 1
            continue
        if ln.startswith("### "):
            doc.add_heading(ln[4:].strip(), level=2)
            i += 1
            continue
        if ln.startswith("## "):
            doc.add_heading(ln[3:].strip(), level=1)
            i += 1
            continue
        if ln.startswith("# "):
            i += 1  # title already rendered on the cover
            continue

        # Blockquote
        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run_text = stripped.lstrip(">").strip()
            add_inline(para, run_text)
            for r in para.runs:
                r.italic = True
            i += 1
            continue

        # Bullets (possibly multi-line continuation)
        if stripped.startswith("- "):
            body = stripped[2:]
            j = i + 1
            while j < len(lines) and lines[j].startswith("  ") and lines[j].strip() and not lines[j].strip().startswith("- "):
                body += " " + lines[j].strip()
                j += 1
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, body)
            i = j
            continue

        # Plain paragraph (join wrapped lines until blank)
        buf = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(("#", "-", ">")):
            buf.append(lines[j].strip())
            j += 1
        para = doc.add_paragraph()
        add_inline(para, " ".join(buf))
        i = j

    # ── Appendix: full screenshot gallery (all captured evidence) ────────────
    doc.add_page_break()
    doc.add_heading("Appendix A — Screenshot gallery (all captured evidence)", level=1)
    module_names = {
        "m1": "Module 1 — India Payroll", "m2": "Module 2 — GST / Invoicing",
        "m3": "Module 3 — Accounting / Journal", "m4": "Module 4 — CRM",
        "m5": "Module 5 — DPDP privacy", "m6": "Module 6 — ITSM",
        "m7": "Module 7 — HR", "m8": "Module 8 — Security / Auth",
    }
    shots = sorted(IMG_DIR.glob("m[0-9]*.png"))
    last_mod = None
    for img in shots:
        prefix = img.name.split("-", 1)[0]
        mod = module_names.get(prefix, prefix)
        if mod != last_mod:
            doc.add_heading(mod, level=2)
            last_mod = mod
        cap = doc.add_paragraph()
        cr = cap.add_run(img.name)
        cr.bold = True
        cr.font.size = Pt(9)
        try:
            doc.add_picture(str(img), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:  # noqa
            doc.add_paragraph(f"[could not embed {img.name}: {e}]")

    doc.save(str(OUT))
    print(f"wrote {OUT}  ({total} findings, {len(findings)} indexed, {len(shots)} gallery images)")


if __name__ == "__main__":
    main()
