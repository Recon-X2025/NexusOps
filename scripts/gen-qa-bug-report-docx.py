#!/usr/bin/env python3
"""Generate the QA verification bug report DOCX from qa-run/findings.json."""
import json
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = "/Users/kathikiyer/Documents/NexusOps"
FINDINGS = os.path.join(ROOT, "qa-run", "findings.json")
SHOTS = os.path.join(ROOT, "qa-run", "screenshots")
OUT = os.path.join(ROOT, "qa-run", "CoheronConnect_QA_Bug_Report_%s.docx" % date.today().isoformat())

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6E, 0x00)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SEV_COLOR = {
    "Critical": RED,
    "High": RED,
    "Medium": AMBER,
    "Low": GREEN,
}

with open(FINDINGS) as f:
    bugs = json.load(f)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, color=None, size=11, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---- Title page ----
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("CoheronConnect")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("QA Verification — Bug Report")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("End-to-end manual QA pass · every module, page and action button\nDate: %s   ·   Environment: Local test database (port 5433)" % date.today().isoformat())
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---- Executive summary ----
heading("Executive summary", size=16)
sev_counts = {}
for b in bugs:
    sev_counts[b["severity"]] = sev_counts.get(b["severity"], 0) + 1
open_count = sum(1 for b in bugs if b.get("status", "").upper() == "OPEN")

para(
    "A deep, manual, module-by-module QA pass was run against the entire CoheronConnect "
    "application. Each of the eight functional modules was driven through a real browser: every "
    "route was loaded, every page inspected, and action buttons (create / save / submit) exercised "
    "end to end with database side-effects verified directly. All testing was performed against the "
    "local test database so no production data was touched.",
)
para(
    "The application is largely stable: core money paths (journal balancing, purchase-order totals, "
    "payroll), create/read/update flows across HR, Finance, Legal, Knowledge Base, Webhooks and API "
    "Keys all behaved correctly, and security-sensitive stores (API-key and webhook secrets, PAN/Aadhaar) "
    "correctly persist only hashes, never raw values. %d defects were identified, summarised below."
    % len(bugs),
)

# summary counts line
counts_bits = []
for sev in ["Critical", "High", "Medium", "Low"]:
    if sev in sev_counts:
        counts_bits.append("%d %s" % (sev_counts[sev], sev))
para("Total findings: %d  (%s)  ·  %d currently OPEN." % (len(bugs), ", ".join(counts_bits), open_count),
     bold=True, color=NAVY)

# ---- Summary table ----
heading("Findings at a glance", size=14)
tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, txt in enumerate(["ID", "Severity", "Module", "Title", "Status"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    shade_cell(hdr[i], "1F3355")

for b in bugs:
    cells = tbl.add_row().cells
    cells[0].text = b["id"]
    # severity cell
    cells[1].text = ""
    sp = cells[1].paragraphs[0]
    sr = sp.add_run(b["severity"])
    sr.bold = True
    sr.font.color.rgb = SEV_COLOR.get(b["severity"], GREY)
    sr.font.size = Pt(10)
    cells[2].text = b["module"]
    cells[3].text = b["title"]
    cells[4].text = b.get("status", "")
    for c in (cells[0], cells[2], cells[3], cells[4]):
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(10)

doc.add_paragraph()

# ---- Detailed findings ----
heading("Detailed findings", size=16)

for b in bugs:
    sev = b["severity"]
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run("%s — %s" % (b["id"], b["title"]))
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    sline = doc.add_paragraph()
    sline.paragraph_format.space_after = Pt(6)
    sr = sline.add_run("Severity: ")
    sr.bold = True
    sr.font.size = Pt(10)
    sv = sline.add_run(sev)
    sv.bold = True
    sv.font.size = Pt(10)
    sv.font.color.rgb = SEV_COLOR.get(sev, GREY)
    mr = sline.add_run("      Module: %s      Page: %s      Status: %s"
                       % (b["module"], b["page"], b.get("status", "")))
    mr.font.size = Pt(10)
    mr.font.color.rgb = GREY

    def field(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lr = p.add_run(label + ": ")
        lr.bold = True
        lr.font.size = Pt(10.5)
        vr = p.add_run(value)
        vr.font.size = Pt(10.5)

    field("Steps to reproduce", b["steps"])
    field("Expected", b["expected"])
    field("Actual", b["actual"])
    field("Root cause", b["rootCause"])
    if b.get("consoleErrors") and b["consoleErrors"].lower() != "none":
        field("Console errors", b["consoleErrors"])

    # screenshot
    shot = b.get("screenshot")
    if shot:
        shot_path = os.path.join(ROOT, "qa-run", shot)
        if os.path.exists(shot_path):
            try:
                doc.add_picture(shot_path, width=Inches(6.0))
                cap = doc.paragraphs[-1]
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                capn = doc.add_paragraph()
                capn.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = capn.add_run("Screenshot: %s" % os.path.basename(shot))
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = GREY
            except Exception as e:
                field("Screenshot", "%s (embed failed: %s)" % (shot, e))

# ---- Coverage appendix ----
doc.add_page_break()
heading("Appendix — coverage", size=16)
para("The pass covered all eight functional modules. For each, routes were swept for load/JS errors, "
     "and representative create/save actions were exercised against the test database and then cleaned up.",
     after=8)

coverage = [
    ("1 · Command Center + Administration", "User management, org settings. BUG-001 found (duplicate invite modal)."),
    ("2 · IT Services", "Service Desk, Change & Release, Field Service, Assets, CMDB — verified."),
    ("3 · Security & Compliance", "SecOps, GRC, ESG, Approvals & Workflow — verified."),
    ("4 · People & Workplace", "HR, Payroll, Recruitment, Performance. BUG-002 found (recruitment deep-links)."),
    ("5 · Customer & Sales", "CSM, CRM, Service Catalog, Surveys — verified."),
    ("6 · Finance & Procurement", "Journal (balance guard OK), Vendors, PO (accrual auto-posts), Contracts. BUG-003 + BUG-004 found."),
    ("7 · Legal & Governance", "Legal delivery, Corporate Secretarial (board directors, MCA/ROC filings) — verified."),
    ("8 · Strategy, Knowledge, Settings, Onboarding", "Strategy/PMO dashboards, Knowledge Base CRUD, Webhooks, API Keys, Integrations, Onboarding wizard — verified."),
]
for name, note in coverage:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(name + " — ")
    r.bold = True
    r.font.size = Pt(10.5)
    n = p.add_run(note)
    n.font.size = Pt(10.5)

doc.save(OUT)
print("Wrote", OUT)
