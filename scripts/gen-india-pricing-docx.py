#!/usr/bin/env python3
"""Generate the India Market Pricing Strategy DOCX (v1.0, revised).

Mirrors the US Pricing Strategy structure. The India book is anchored to the
fragmented Indian SMB stack (helpdesk + PM + HRMS + ops/tail apps), uses a flat
per-org, size-banded model with the whole platform included in every tier (never
per-seat, never per-module), and includes founding-partner (pilot) terms that
mirror the "100 Believers" motion.

Revised ladder decided in this session. Billing is FLAT PER-ORG (not per-seat);
the per-seat figures are implied economics only, tuned so the implied per-seat cost
falls monotonically from Rs 800 (Start) to Rs 500 (Custom) as bands grow -- no
inversion:
  Start   0-10 emp   app free + Rs 8,000/mo platform fee (self-serve; ~Rs 800/seat at cap)
  Launch  10-40 emp  Rs 28,000 flat   (~Rs 700/seat at 40)
  Grow    40-75 emp  Rs 45,000 flat   (~Rs 600/seat at 75)
  Scale   75-250 emp Rs 1,37,500 flat (~Rs 550/seat at 250)
  Custom  250+ emp   custom quote     (~Rs 500/seat basis)

All rupee figures are proposed/working numbers, not recovered from any prior file.
Stack-anchor ranges are market estimates to be validated against a real prospect's
invoice during discovery. The India book is independent of the US book and is never
cross-referenced on customer-facing surfaces.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/kathikiyer/Documents/NexusOps/docs-word/CoheronConnect_India_Pricing_Strategy_2026-07-13.docx"

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

RS = "\u20b9"


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, color=None, size=11, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def table(headers, rows, bold_last=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            r = cells[i].paragraphs[0].add_run(val)
            r.font.size = Pt(9.5)
            if bold_last and ridx == len(rows) - 1:
                r.bold = True
    return t


# ---- Title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("CoheronConnect")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
r = st.add_run("India Market Pricing Strategy")
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
r = meta.add_run("Version 1.0 (revised)  •  13 July 2026  •  Internal — Confidential")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---- Governing principle ----
heading("The governing principle: anchored to the stack we replace", 14)
para(
    "The India book is anchored to what an Indian SMB already pays for its fragmented stack "
    "(helpdesk/ITSM + project management + HRMS/payroll + ops and tail apps), not to a "
    "converted foreign price. It is an independent price book: it is never derived from, nor "
    "cross-referenced against, the US book on any customer-facing surface. Same tiers, same "
    "packaging as the US — only the numbers and the local anchor differ."
)

# ---- 1. Model ----
heading("1. The India Pricing Model", 15)
para(
    "Flat platform pricing per tier — not per-seat, never per-module. Every tier includes the "
    "whole platform; price scales by company stage and size. The application itself is free — "
    "there is no software licence cost. What customers pay for is the platform (hosting, "
    "operations, security, uptime) and, at higher tiers, the human support and SLAs around it. "
    "\u201cFlat price, whole company on the platform\u201d reinforces the operating-system "
    "positioning and avoids per-seat fatigue as headcount grows."
)
table(
    ["Tier", "Size band", "Flat price (per org)", "Implied /seat at band top", "Support"],
    [
        ("Start", "0-10 emp", f"App free + {RS}8,000/mo platform fee",
         f"~{RS}800", "Self-serve only"),
        ("Launch", "10-40 emp", f"{RS}28,000/mo flat",
         f"~{RS}700", "Self-serve + email"),
        ("Grow", "40-75 emp", f"{RS}45,000/mo flat",
         f"~{RS}600", "Guided onboarding + priority"),
        ("Scale", "75-250 emp", f"{RS}1,37,500/mo flat",
         f"~{RS}550", "Dedicated CS + SLA"),
        ("Scale (Custom)", "250+ emp", "Custom quote",
         f"~{RS}500", "Named CSM + custom SLA"),
    ],
)
para(
    "Billing is flat per organisation — a Launch customer pays "
    f"{RS}28,000 whether they are 25 or 40 people. The \u201cimplied per-seat\u201d column is "
    "economics, not a billing basis: prices are tuned so the implied per-seat cost falls "
    f"monotonically from ~{RS}800 (Start) to ~{RS}500 (Custom) as companies grow. This removes "
    "the earlier inversion where a smaller band was cheaper per head than a larger one. A minor, "
    "standard flat-band edge effect remains just below each band top; it is not eliminable without "
    "per-seat billing, which is deliberately not used.",
    size=10,
)
para(
    "Start is the lead-entry tier. The application is free to use; a flat "
    f"{RS}8,000/mo platform fee covers the infrastructure and platform operations the tenant "
    "consumes. Start is self-serve only — documentation, in-app onboarding, and ticket/community "
    "support, but no white-glove: no dedicated success manager, no hands-on implementation, no SLA. "
    "That absence is deliberate — human support is a tier differentiator, and the need for it is a "
    "natural trigger to move up to Launch or Grow.",
    size=10,
)
para(
    "Size bands prevent coarse-band underpricing: without ceilings, a 10-person and a 60-person "
    "company would pay the same. Start tops out at 10 employees, Launch at 40, Grow at 75; above "
    "that a company moves to Scale, and above ~250 employees it is quoted Custom rather than "
    "capped at the list ceiling.",
    size=10,
)

# ---- 2. Why the prices hold ----
heading("2. Why These Prices Hold Against the Market", 15)
para(
    "The comparison an Indian buyer actually runs is not against another all-in-one — it is "
    "against the stack they already pay for. A typical 50-200 person Indian company runs a "
    "helpdesk/ITSM tool, a project tool, an HRMS/payroll system, plus a tail of smaller "
    "subscriptions:"
)
table(
    ["Stack component", "Typical India spend /mo", "CoheronConnect"],
    [
        ("Helpdesk / ITSM (Freshservice-class)", f"{RS}25,000 - 60,000", "Included"),
        ("Work / project management", f"{RS}15,000 - 35,000", "Included"),
        ("HRMS / payroll (Keka / greytHR-class)", f"{RS}20,000 - 50,000", "Included"),
        ("Ops database, approvals, misc. tail apps", f"{RS}15,000 - 40,000", "Included"),
        ("Typical stack total", f"{RS}75,000 - 1,85,000", f"{RS}35,000 (Grow) flat"),
    ],
    bold_last=True,
)
para(
    f"At {RS}45,000 flat, Grow sits well below the {RS}75,000-1,85,000 stack it replaces — an "
    "easy switch on cost alone, while still credible and far above a point-tool price. Pricing to a "
    "converted foreign number, or to bare per-seat economics, would leave most of this "
    "stack-replacement value on the table and give no room to fund support at the higher tiers.",
)

# ---- 3. Price-book hygiene ----
heading("3. Price-Book Hygiene", 15)
para(
    "Independent, geo-separated pricing pages. Indian visitors see INR pricing only; the book is "
    "never presented alongside, or derived from, any other market's pricing. The cross-market "
    "comparison is internally understood and externally invisible.",
    size=10,
)
para(
    "Annual incentive. Monthly billing is common in the Indian SMB market; annual commitments are "
    "incentivised (a fixed number of months free or an equivalent discount) to stabilise revenue.",
    size=10,
)
para(
    "Negotiation posture. List is list. Standard concessions on annual commitments are bounded; "
    "anything deeper is founder-approved and exchanged for term length, case-study rights, or "
    "reference calls — never given away unpriced.",
    size=10,
)
para(
    "Cross-border edge case. An Indian company transacting through a foreign entity buys on the "
    "book of the entity that signs the contract and receives the invoice.",
    size=10,
)

# ---- 4. Founding-partner terms ----
heading("4. Founding-Partner (Pilot) Terms", 15)
para(
    "Early Indian customers should be priced off list, not around it — the discount must be framed "
    "as a named program that protects the anchor rather than eroding it. This mirrors the "
    "\u201c100 Believers\u201d founding-customer motion."
)
table(
    ["Element", "Recommendation"],
    [
        ("Frame", "\u201cFounding partner\u201d — a named program, not a discount."),
        ("Price", f"{RS}27,000/mo (~40% off the {RS}45,000 Grow list) on a 12-month agreement; "
         "list price stated on the order form with the program credit shown against it."),
        ("Exchange", "Feedback sessions, India case study + logo rights, 2 reference calls, "
         "input on the roadmap."),
        ("Expiry", "Program pricing holds 24 months, then steps to then-current list minus a "
         "loyalty tier — stated in the agreement so there is no renewal shock."),
        ("Scarcity", "A capped founding-partner cohort. Scarcity keeps the concession from "
         "becoming the street price."),
    ],
)

# ---- 5. The Rs 15 Cr ARR view ----
heading(f"5. The {RS}15 Cr ARR View on This Ladder", 15)
para(
    f"{RS}15 Cr ARR = {RS}1.25 Cr MRR. Start is counted as revenue — its "
    f"{RS}8,000/mo platform fee is real recurring revenue and contributes to ARR, even though "
    "it is largely cost-recovery and therefore margin-light. It is revenue for top-line purposes "
    "and a low-margin line for profitability purposes; both are true and are kept distinct.",
)
table(
    ["Tier", "Price /mo", "Illustrative customers", "MRR", "ARR"],
    [
        ("Start", f"{RS}8,000", "200", f"{RS}16.0 L", f"{RS}1.92 Cr"),
        ("Launch", f"{RS}28,000", "130", f"{RS}36.4 L", f"{RS}4.37 Cr"),
        ("Grow", f"{RS}45,000", "70", f"{RS}31.5 L", f"{RS}3.78 Cr"),
        ("Scale", f"{RS}1,37,500", "20", f"{RS}27.5 L", f"{RS}3.30 Cr"),
        ("Scale (Custom)", f"{RS}2,00,000", "7", f"{RS}14.0 L", f"{RS}1.68 Cr"),
        ("Total", "—", "~427", f"{RS}1.25 Cr", f"~{RS}15.05 Cr"),
    ],
    bold_last=True,
)
para(
    f"Roughly 427 paying customers reach ~{RS}15 Cr ARR on this ladder. Start contributes "
    f"~{RS}1.92 Cr (~13% of ARR) — meaningful top-line, but margin-light; the profit engine is "
    "Launch, Grow, Scale and Custom. Counting Start as revenue (rather than excluding it) is the "
    "honest top-line view.",
    size=10,
)

# ---- 6. Open items ----
heading("6. Open Items", 15)
para("1. Validate the stack-basket figures against a real prospect's actual current spend during "
     "discovery — their real invoice total is the best sales asset available.", size=10, after=2)
para("2. Confirm what the Start platform fee covers in writing (infra + platform ops + self-serve "
     "support) and the usage/headcount limits that trigger a move to Launch.", size=10, after=2)
para("3. Lock the exact Launch-to-Grow, Grow-to-Scale and Scale-to-Custom thresholds in billing "
     "(10 / 40 / 75 / 250 employees are the working defaults).", size=10, after=2)
para(f"4. Confirm the Custom planning basis ({RS}2,00,000/mo used in the ARR view) and the "
     f"founding-partner cohort cap.", size=10)

# ---- Bottom line ----
heading("Bottom line", 15, GREEN)
para(
    f"India list (flat per-org, whole platform included, tuned so implied per-seat falls "
    f"{RS}800\u2192{RS}500): Start — app free + {RS}8,000/mo platform fee, self-serve only; "
    f"Launch {RS}28,000 (10-40 emp); Grow {RS}45,000 (40-75 emp); "
    f"Scale {RS}1,37,500 (75-250); Custom above 250. Priced by stage and size, never by module "
    f"or seat. Anchored to the {RS}75,000-1,85,000 fragmented stack it replaces and independent of "
    f"any other market's book. Start counts as revenue (~{RS}1.92 Cr, margin-light); ~427 "
    f"customers reach ~{RS}15 Cr ARR. Founding partners land at {RS}27,000/mo. All figures are "
    "working numbers pending validation against real prospect spend.",
    italic=True, color=GREY, size=10,
)

doc.save(OUT)
print(f"Wrote {OUT}")
