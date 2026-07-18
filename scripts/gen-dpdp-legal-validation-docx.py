#!/usr/bin/env python3
"""Generate the DPDP Erasure — Legal Validation Request as a circulatable DOCX.

Mirrors the styling conventions in gen-gap-docx.py. Source of truth is
docs/DPDP_ERASURE_LEGAL_VALIDATION.md; keep them in sync if either changes.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/Users/kathikiyer/Documents/NexusOps/docs-word/CoheronConnect_DPDP_Erasure_Legal_Validation_2026-07-16.docx"

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6E, 0x00)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, size=15, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, bold=False, color=None, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(text, color=ACCENT):
    """Indented emphasis paragraph for 'counsel to confirm' prompts."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = color
    return p


def set_cell(cell, text, bold=False, color=None, size=10):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def header_row(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        # navy shading
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F3355")
        tcPr.append(shd)


# ── Title block ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("DPDP Erasure — Legal Validation Request")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

meta = [
    ("Prepared for:", "Indian privacy counsel"),
    ("Prepared by:", "CoheronConnect — Product / Compliance"),
    ("Date:", "2026-07-16"),
    ("Decision requested by:", "_______________"),
    ("Product:", "CoheronConnect — multi-tenant Enterprise Operations Platform "
     "(payroll, tax, accounting, procurement, CRM, secretarial/ROC). Data resident in India."),
]
for k, v in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rk = p.add_run(k + " ")
    rk.bold = True
    rk.font.size = Pt(11)
    rv = p.add_run(v)
    rv.font.size = Pt(11)

# Purpose banner
para(
    "Purpose of this document. We are wiring the automated fulfilment of Data Principal "
    "erasure requests (DPDP §12). Because the platform also holds financial records under "
    "statutory retention (RBI / Companies Act 2013 / Income Tax Act), we need counsel to "
    "validate and sign off the rules below before any destructive automation is enabled. "
    "Nothing destructive runs until this is signed — the erasure engine currently operates "
    "in dry-run (log-only) mode.",
    italic=True, color=GREY, after=6,
)
para(
    "What we need from counsel: (a) confirm the legal reconciliation in §1, (b) fill and "
    "sign the column-by-column erasure map in §4, (c) answer the specific questions in §5.",
    italic=True, color=GREY, after=10,
)

# ── §1 ───────────────────────────────────────────────────────────────────────
heading("1. The core legal question")
para("Can we honour a DPDP erasure request without deleting financial records we are "
     "legally required to retain?", bold=True)
para("Our proposed reconciliation:")
bullet(" grants the Data Principal a right to erasure of personal data.", bold_lead="DPDP §12")
bullet(" (our understanding) does not require erasure where retention is necessary for "
       "compliance with any law in force.", bold_lead="DPDP §8(7)")
bullet("For financial records we propose to retain the financial fact and figures (as "
       "statute requires) while severing / anonymising the personal-identity link "
       "(satisfying the erasure right). We do not hard-delete statutory financial records.")
callout("➡ Counsel to confirm: Is anonymisation / de-identification of the personal element "
        "(while retaining the underlying financial record for the statutory period) a valid "
        "means of satisfying a DPDP §12 erasure request?   ☐ Confirmed   ☐ Needs revision")

# ── §2 ───────────────────────────────────────────────────────────────────────
heading("2. Proposed retention floor")
para("We propose a blanket 8-year retention floor on financial records (payslips, journal "
     "entries, invoices, purchase orders, tax filings), after which the identity link may "
     "be anonymised on request. The record is held across storage tiers (hot / warm / cold) "
     "for the full 8 years.")
callout("➡ Counsel to confirm: Is a uniform 8-year floor correct, or do specific record "
        "classes require different periods (e.g., payroll vs. GST vs. statutory registers "
        "vs. audit)?   ☐ 8y blanket confirmed   ☐ Per-class periods required (specify in §4)")

# ── §3 ───────────────────────────────────────────────────────────────────────
heading("3. Categories of personal data we hold (plain-language)")
t3 = doc.add_table(rows=1, cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.LEFT
t3.style = "Table Grid"
header_row(t3, ["Category", "Examples in the system", "Statutory retention?"])
rows3 = [
    ("Employee identity & financial", "Name, email, PAN (tax ID), Aadhaar, bank account, address", "Yes — payroll/tax records"),
    ("Director identity", "Name, DIN, PAN, Aadhaar, residential address", "Yes — statutory registers"),
    ("Vendor / supplier contacts", "Contact person name, email, phone; GSTIN, PAN", "Entity: yes; person contact: no"),
    ("Financial transaction records", "Payslips, journal entries, invoices, POs (amounts, dates)", "Yes — 8 years"),
    ("CRM contacts", "Prospect first/last name, email, phone", "No"),
    ("Audit trail", "Tamper-evident log of who-changed-what", "Yes — integrity/security"),
]
for a, b, c in rows3:
    cells = t3.add_row().cells
    set_cell(cells[0], a, bold=True)
    set_cell(cells[1], b)
    set_cell(cells[2], c)

# ── §4 ───────────────────────────────────────────────────────────────────────
heading("4. Erasure map — for counsel to complete and sign")
para("For each personal-data field, counsel confirms the Action and Retention on an "
     "erasure request. Actions:")
bullet(" — remove the value/row entirely (no retention obligation).", bold_lead="Delete")
bullet(" — retain the record; overwrite the personal field with a redaction marker; where "
       "a non-reversible reference is needed for statutory audit, keep a one-way hash.", bold_lead="Anonymise")
bullet(" — never store the raw value at all (masked hash only), so there is nothing to "
       "erase later.", bold_lead="Minimise-at-source")
bullet(" — must be kept for the stated statutory period; not erasable until it expires.", bold_lead="Retain-with-reason")
bullet(" — for the tamper-evident audit log, which cannot be edited in place without "
       "breaking its integrity chain; a separate \u201credacted for DPDP\u201d event is recorded "
       "instead.", bold_lead="Redaction side-event")

t4 = doc.add_table(rows=1, cols=8)
t4.alignment = WD_TABLE_ALIGNMENT.LEFT
t4.style = "Table Grid"
header_row(t4, ["#", "Data field", "Category", "Fin?", "Proposed action", "Statutory basis", "Retain until", "Sign-off"])
rows4 = [
    ("1", "Employee name / email", "Identity", "Yes", "Anonymise", "RBI / IT Act", "pay + 8y", "☐ approve ☐ amend"),
    ("2", "Employee PAN (tax ID)", "Tax ID", "Yes", "Anonymise; keep one-way hash", "IT Act / RBI", "pay + 8y", "☐ approve ☐ amend"),
    ("3", "Employee Aadhaar", "Gov ID", "Yes", "Minimise-at-source (masked hash; never raw)", "DPDP §8 / Aadhaar Act", "n/a (never raw)", "☐ approve ☐ amend"),
    ("4", "Employee bank account / IFSC", "Financial", "Yes", "Anonymise", "RBI", "pay + 8y", "☐ approve ☐ amend"),
    ("5", "Employee address / city / state", "Contact", "Yes", "Anonymise", "RBI / IT Act", "pay + 8y", "☐ approve ☐ amend"),
    ("6", "Director Aadhaar", "Gov ID", "Yes", "Minimise-at-source (masked hash; never raw)", "DPDP §8 / Aadhaar Act", "n/a (never raw)", "☐ approve ☐ amend"),
    ("7", "Director PAN / residential address", "Identity", "Yes", "Anonymise", "Companies Act", "cessation + 8y", "☐ approve ☐ amend"),
    ("8", "Payslip (amounts + tax)", "Financial", "Yes", "Retain-with-reason (identity via hash link)", "RBI / IT Act", "pay + 8y", "☐ approve ☐ amend"),
    ("9", "Journal entry (amounts, actor FK)", "Financial", "Yes", "Retain-with-reason (actor link nulled)", "Companies Act", "post + 8y", "☐ approve ☐ amend"),
    ("10", "Invoice / PO (GSTIN, amounts)", "Financial", "Yes", "Retain-with-reason", "GST / IT Act", "invoice + 8y", "☐ approve ☐ amend"),
    ("11", "Vendor contact person name/email/phone", "Contact", "No", "Delete / anonymise (vendor entity retained)", "—", "on erasure", "☐ approve ☐ amend"),
    ("12", "Vendor GSTIN / PAN (entity)", "Tax ID", "Yes", "Retain-with-reason", "GST / IT Act", "relationship + 8y", "☐ approve ☐ amend"),
    ("13", "CRM contact name / email / phone", "Contact", "No", "Delete", "—", "on erasure", "☐ approve ☐ amend"),
    ("14", "Audit-log entry referencing the Principal", "Mixed", "—", "Redaction side-event (chain preserved)", "tamper-evidence", "chain life", "☐ approve ☐ amend"),
]
for row in rows4:
    cells = t4.add_row().cells
    for i, val in enumerate(row):
        set_cell(cells[i], val, size=9)
para("Add rows for any field counsel believes is missing.", italic=True, color=GREY, after=8)

# ── §5 ───────────────────────────────────────────────────────────────────────
heading("5. Specific questions for counsel")
q = [
    ("Anonymisation as erasure (§1).", " Does retaining a financial record while anonymising "
     "the personal element satisfy a DPDP §12 erasure request? Any documentation/evidence we "
     "must generate to demonstrate compliance?"),
    ("Retention period (§2).", " Confirm 8-year blanket vs. per-class. Which record classes, "
     "if any, require longer (e.g., statutory registers, litigation hold)?"),
    ("Aadhaar minimisation (rows 3, 6).", " We propose to never store raw Aadhaar — masked "
     "last-4 + one-way hash only (mirroring our existing e-sign handling). Is there any lawful "
     "purpose in payroll/secretarial that requires retrievable full Aadhaar? If not, we "
     "minimise at onboarding and backfill existing records."),
    ("PAN as retained hash (row 2).", " Is keeping a one-way hash of PAN (for statutory audit "
     "matching) acceptable after anonymising the raw PAN, or must raw PAN be retained in full "
     "for the 8-year window?"),
    ("Erasure inside the retention window.", " When a Principal requests erasure but their "
     "financial records are still inside the 8-year window, we propose to anonymise the "
     "contact/identity fields now but retain the figures + statutory identifiers until the "
     "window expires, then anonymise those too. Confirm this staged approach is acceptable, and "
     "what we must tell the Principal (statutory-retention notice wording)."),
    ("Audit-trail redaction (row 14).", " Our security audit log is tamper-evident (a "
     "cryptographic chain); editing an entry breaks that chain and is itself a compliance "
     "feature. We propose to record a separate \u201credacted for DPDP\u201d event rather than alter the "
     "original entry. Is preserving the immutable entry (with a redaction marker recorded "
     "alongside) acceptable, or must the underlying PII in the log be destroyed?"),
    ("Crypto-shredding (future).", " Once per-subject key encryption is in place, \u201cerasure\u201d "
     "can be achieved by destroying the encryption key — the record remains but is permanently "
     "unreadable. Do you accept key-destruction as satisfying erasure?"),
    ("Data Principal notice.", " What statutory wording must we return to a Principal whose "
     "erasure is partially deferred due to retention (i.e., \u201cyour contact data has been "
     "erased; your financial records are retained under [statute] until [date]\u201d)?"),
]
for i, (lead, body) in enumerate(q, 1):
    p = doc.add_paragraph(style="List Number")
    rl = p.add_run(lead)
    rl.bold = True
    p.add_run(body)

# ── §6 ───────────────────────────────────────────────────────────────────────
heading("6. What happens after sign-off")
para("Once counsel completes §4 and answers §5:")
for i, s in enumerate([
    "Engineering encodes the signed map into the erasure engine (currently dry-run only).",
    "A retention guard is added so no record is anonymised/deleted before its \u201cretain until\u201d date.",
    "Aadhaar minimisation is applied at onboarding and backfilled.",
    "Per-category test erasure requests are run to prove: retained records survive with "
    "figures intact, identity is severed, and the audit chain still verifies.",
    "Only then is destructive erasure switched on in production.",
], 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)
para("No destructive erasure runs in production until this document is signed.",
     bold=True, color=RED, after=12)

# ── Sign-off block ───────────────────────────────────────────────────────────
heading("Counsel sign-off")
para("Name: _______________     Firm: _______________", after=10)
para("Signature: _______________     Date: _______________", after=12)
para("This document requests legal validation only. It does not itself change any system "
     "behaviour.", italic=True, color=GREY)

doc.save(OUT)
print("wrote", OUT)
