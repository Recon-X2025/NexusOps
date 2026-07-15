#!/usr/bin/env python3
"""Generate the website security & compliance brief DOCX.

Non-technical, circulatable to the website/marketing team. States ONLY controls
that exist in the code today (each verified file:line in the internal audit),
separates "Available Today" from "On the Roadmap" from "Do NOT publish", and
gives ready-to-paste, legally-safe wording. No certification claims.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/kathikiyer/Documents/NexusOps/docs-word/CoheronConnect_Security_Compliance_Website_Brief_2026-07-13.docx"

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6E, 0x00)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, color=None, size=11, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def status_run(cell, value):
    p = cell.paragraphs[0]
    r = p.add_run(value)
    r.bold = True
    if value.startswith("LIVE"):
        r.font.color.rgb = GREEN
    elif value.startswith("PARTIAL"):
        r.font.color.rgb = AMBER
    else:
        r.font.color.rgb = RED


def table(headers, rows, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if status_col is not None and i == status_col:
                status_run(cells[i], val)
            else:
                r = cells[i].paragraphs[0].add_run(val)
                r.font.size = Pt(10)
    return t


# ---- Title block ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("CoheronConnect")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
r = st.add_run("Security & Compliance — Website Messaging Brief")
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
r = meta.add_run(
    "Prepared 13 July 2026  •  For the website / marketing team  •  "
    "Every 'Available Today' item below is implemented in the live product and was "
    "verified against the source code. Nothing here is aspirational."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---- Golden rule ----
heading("The one rule for this page", 15)
para(
    "Describe the security controls we actually operate. Never claim a certification, "
    "audit, or attestation we have not earned, and never imply we are \"compliant\" with a "
    "law as a property of the software. Controls are facts we can defend; certifications "
    "and \"compliance\" are formal statuses we do not yet hold. Getting this wrong is a "
    "legal and reputational risk, not just a wording preference.",
    bold=False,
)

# ---- Available today ----
heading("1. Available today — safe to publish", 15)
para(
    "These controls are built into the product now. Suggested public wording is in the "
    "right-hand column — plain, factual, no over-claiming.",
    italic=True, color=GREY, size=9.5,
)

table(
    ["Control (internal)", "Status", "Wording you can publish"],
    [
        ("Password security", "LIVE",
         "Credentials are hashed with bcrypt (industry-standard, work factor 12) — never stored in plain text."),
        ("Session protection", "LIVE",
         "Session tokens are hashed at rest with configurable expiry and server-side revocation."),
        ("Role-based access control", "LIVE",
         "Granular role-based access control across every module, enforced on every request."),
        ("Tamper-evident audit log", "LIVE",
         "A tamper-evident audit trail (cryptographic hash chain) records changes and can be verified for integrity."),
        ("Rate limiting / abuse protection", "LIVE",
         "Brute-force and abuse protection with request rate limiting on authentication and APIs."),
        ("Input validation", "LIVE",
         "Strict input validation on all requests; database access is fully parameterised (no string-built SQL)."),
        ("Encrypted secrets", "LIVE",
         "Third-party integration credentials are encrypted (AES-256) before storage."),
        ("Multi-tenant isolation", "LIVE",
         "Every customer's data is isolated per organisation and scoped on every query."),
        ("Security headers & CORS", "LIVE",
         "Hardened HTTP security headers and strict cross-origin controls."),
        ("Step-up re-authentication", "LIVE",
         "Sensitive actions can require step-up re-authentication for higher-risk roles."),
        ("Privacy request tooling (DPDP-aligned)", "LIVE",
         "Built-in tooling to manage consent, data-subject requests, and breach records with statutory response tracking."),
    ],
    status_col=1,
)

doc.add_paragraph()
para(
    "Suggested section headline for the site: \"Security built in, not bolted on.\" "
    "Suggested sub-line: \"Role-based access, tamper-evident audit logging, encrypted "
    "credentials, and abuse protection — on a single, multi-tenant platform.\"",
    italic=True, color=GREY, size=9.5,
)

# ---- Roadmap ----
heading("2. On the roadmap — publish only as forward-looking", 15)
para(
    "These are partially built or planned. If mentioned at all, they MUST be framed as "
    "future/coming — use \"on our roadmap\", \"coming soon\", or \"in development\". Do not "
    "present them as available.",
    italic=True, color=GREY, size=9.5,
)
table(
    ["Item", "Status", "If published, say…"],
    [
        ("Two-factor authentication (TOTP app)", "PARTIAL",
         "\"Two-factor authentication is on our near-term roadmap.\""),
        ("Encryption of personal data fields at rest", "PARTIAL",
         "\"Field-level encryption for sensitive data is in development.\""),
        ("Managed key service (KMS) integration", "PARTIAL",
         "\"Customer-managed encryption keys are planned.\""),
        ("Database row-level security", "PARTIAL",
         "Internal defence-in-depth — best not mentioned publicly at all yet."),
        ("Automated notice delivery (breach / DSR)", "PARTIAL",
         "Internal — do not describe notifications as automatically delivered yet."),
        ("Data erasure execution across all modules", "PARTIAL",
         "Internal — see the DPDP readiness checklist; not a public claim yet."),
    ],
    status_col=1,
)

# ---- Do NOT publish ----
heading("3. Do NOT publish — until formally earned", 15)
para(
    "The following must stay OFF the website until we hold the actual status. Using them "
    "prematurely is a misrepresentation.",
    italic=True, color=GREY, size=9.5,
)
bullet(" — we are not audited/attested. Remove entirely until we complete an audit.", "\"SOC 2 Type II\"")
bullet(" — we are not certified. Remove entirely until certified.", "\"ISO 27001 certified\"")
bullet(" — compliance is an organisational status, not a software feature. Say \"DPDP-aligned tooling\" instead.", "\"DPDP compliant\" / \"fully compliant\"")
bullet(" (unqualified) — we do not yet encrypt all personal-data fields at rest. Only claim encryption where it is true (secrets, transport).", "\"End-to-end encrypted\" / \"encrypted at rest\"")
bullet(" — no penetration test or bug-bounty is in place to reference.", "\"Bank-grade\" / \"military-grade\" / \"pen-tested\"")

# ---- Safe phrasing bank ----
heading("4. Ready-to-paste safe phrases", 15)
bullet("\"Security is designed into CoheronConnect from the ground up.\"")
bullet("\"Role-based access control governs who can see and do what, on every request.\"")
bullet("\"A tamper-evident audit trail records activity for accountability and integrity checks.\"")
bullet("\"Credentials are hashed with industry-standard algorithms and never stored in plain text.\"")
bullet("\"Built-in privacy tooling helps you manage consent and respond to data-subject requests, "
       "with statutory response tracking — designed to support your DPDP obligations.\"")
bullet("\"We are actively investing in our security roadmap, including two-factor authentication and "
       "field-level encryption.\"")

# ---- Footer note ----
doc.add_paragraph()
para(
    "Note: This brief lists controls, not legal advice. Any compliance or certification "
    "language on the public site should be reviewed by counsel before publishing. A "
    "companion \"DPDP Readiness Checklist\" tracks the legal/organisational items that "
    "sit outside the software.",
    italic=True, color=GREY, size=9,
)

doc.save(OUT)
print(f"Wrote {OUT}")
