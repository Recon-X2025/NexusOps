#!/usr/bin/env python3
"""Generate the US Market Pricing Strategy DOCX (v2.0).

v2.0 re-cuts the US ladder to 5 bands (100% tier-matched to the India book) and
fixes a per-seat inversion in the v1.1 numbers. Billing stays FLAT PER-ORG; the
per-seat figures are implied economics only, tuned so implied per-seat falls
monotonically as bands grow -- no inversion.

Anchoring: each band is set against US market reality (the fragmented stack the
buyer already pays), then priced at a deliberate ~25-35% discount to that stack --
cheap enough to be an easy switch, not so cheap as to invite quality doubt.

Revised ladder:
  Start   0-10 emp   app free + $99/mo platform fee (self-serve; implied ~$99/seat floor)
  Launch  10-40 emp  $1,600 flat   (~$40/seat at 40)
  Grow    40-75 emp  $2,499 flat   (~$33/seat at 75)
  Scale   75-250 emp $7,000 flat   (~$28/seat at 250)
  Custom  250+ emp   custom quote  (~$24/seat basis)

Note: v2.0 supersedes the v1.1 $1,999 Grow anchor. The design-partner pilot moves to
~$1,499 (40% off the new $2,499 Grow list). US book only; the India book is never
cross-referenced on customer-facing surfaces. All figures are working numbers pending
validation against a real prospect's invoice.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/kathikiyer/Documents/NexusOps/docs-word/CoheronConnect_US_Pricing_Strategy_2026-07-13.docx"

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, color=None, size=11, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def table(headers, rows, bold_last=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            r = cells[i].paragraphs[0].add_run(val)
            r.font.size = Pt(9.5)
            if bold_last and ridx == len(rows) - 1:
                r.bold = True
    return t


# ---- Title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("CoheronConnect")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
r = st.add_run("US Market Pricing Strategy")
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
r = meta.add_run("Version 2.0  •  13 July 2026  •  Internal — Confidential")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---- Governing principle ----
heading("The governing principle: two independent price books", 14)
para(
    "Neither price book derives from the other. Each is anchored to what that market's "
    "buyer already pays. The US book is set against the fragmented US SMB stack "
    "(helpdesk + project management + database/ops tools + tail apps, typically "
    "$1,700-3,700/month all-in). The India book stays anchored to its own local reference "
    "spend. Same tiers, same names, same packaging — only the numbers differ. This is not "
    "a discount or a premium; it is two independent anchors."
)

# ---- 1. Model ----
heading("1. The US Pricing Model", 15)
para(
    "Flat platform pricing per tier — not per-seat. Every competitor in the US reference "
    "basket charges per seat, and per-seat fatigue is a documented pain point for growing "
    "US businesses: costs scale linearly with every hire. \u201cFlat price, whole company on "
    "the platform\u201d is therefore a differentiator in its own right, not just a pricing "
    "choice, and it reinforces the operating-system positioning. Every tier includes the "
    "full platform; price scales by company stage and size, never by module."
)
table(
    ["Tier", "Size band", "Flat price (per org)", "Implied /seat at band top", "Billing"],
    [
        ("Start", "0-10 emp", "App free; $99/mo platform fee",
         "~$99", "Monthly"),
        ("Launch", "10-40 emp", "$1,600/mo flat",
         "~$40", "Annual default"),
        ("Grow", "40-75 emp", "$2,499/mo flat",
         "~$33", "Annual default"),
        ("Scale", "75-250 emp", "$7,000/mo flat",
         "~$28", "Annual contract"),
        ("Scale (Custom)", "250+ emp", "Custom quote",
         "~$24", "Annual contract"),
    ],
)
para(
    "Billing is flat per organisation — a Grow customer pays $2,499 whether they are 55 or 75 "
    "people. The \u201cimplied per-seat\u201d column is economics, not a billing basis: prices "
    "are tuned so implied per-seat cost falls monotonically as companies grow, removing the "
    "inversion in the earlier ladder (where the $1,999 Grow list implied a per-seat rate below "
    "the Scale floor, penalising a company that grew past 75). A minor flat-band edge effect "
    "remains just below each band top; it is not eliminable without per-seat billing, which is "
    "deliberately not used.",
    size=10,
)
para(
    "The ladder is 100% tier-matched to the India book (same five bands and thresholds); only "
    "the numbers and the local anchor differ. Every tier includes the full platform, priced by "
    "stage and size, never by module.",
    size=10,
)
para(
    "Start is free to use — the application carries no licence cost. A flat $99/mo platform "
    "fee covers hosting and platform operations; it is a platform fee, not a paid plan, and it is "
    "self-serve only. This keeps the entry genuinely product-led while establishing a billing "
    "relationship from day one; the need for human support is a natural trigger to move up to "
    "Launch or Grow.",
    size=10,
)
para(
    "Each band is anchored to US market reality — the fragmented stack the buyer already pays — "
    "then priced at a deliberate ~25-35% discount to that stack. Cheap enough to be an obvious "
    "switch, expensive enough to stay credible: a converted-from-abroad price (~$20-25/user) would "
    "seem implausibly cheap and invite quality doubts.",
    size=10,
)

# ---- 2. Why these prices hold ----
heading("2. Why These Prices Hold Against the Market", 15)
para(
    "The comparison a US buyer actually runs is not against another all-in-one — it is against "
    "the stack they already pay for. A typical 20-50 person US company runs a helpdesk, a "
    "project tool, a flexible database, plus a tail of smaller subscriptions:"
)
table(
    ["Stack component", "Typical US spend /mo", "CoheronConnect"],
    [
        ("Helpdesk / support (Zendesk-class)", "$550 - 1,150", "Included"),
        ("Work / project management (Asana-class)", "$300 - 750", "Included"),
        ("Ops database / internal tools (Airtable-class)", "$400 - 900", "Included"),
        ("HR, approvals, misc. tail apps", "$450 - 900", "Included"),
        ("Typical stack total", "$1,700 - 3,700", "$2,499 (Grow) flat"),
    ],
    bold_last=True,
)
para(
    "At $2,499 flat, Grow sits at the low-middle of the stack it replaces — an easy switch on "
    "cost alone, while still credible. Every band is set ~25-35% below the equivalent stack "
    "spend at that size, so the discount is real but never implausible. A converted-from-abroad "
    "price (~$20-25/user) would position the product implausibly cheap for a US buyer and invite "
    "quality doubts; underpricing damages credibility as much as overpricing."
)

# ---- 3. Price-book hygiene ----
heading("3. Price-Book Hygiene", 15)
para(
    "Geo-separated pricing pages. The USD page is never one click from the local page. US "
    "visitors see US pricing only. The cross-market comparison is never volunteered — internally "
    "understood, externally invisible.",
    size=10,
)
para(
    "Annual-first billing. Annual is the default presentation in the US with a 15-20% incentive "
    "versus monthly. This matches US SMB buying norms and stabilises early revenue.",
    size=10,
)
para(
    "Negotiation posture. List is list. Sales can concede 15-25% on annual commitments without "
    "approval; anything deeper is founder-approved and exchanged for term length, case-study "
    "rights, or reference calls — never given away unpriced.",
    size=10,
)
para(
    "Cross-border edge case. An Indian company with US operations buys on the book of the entity "
    "that signs the contract and receives the invoice. The US subsidiary contracts on US paper at "
    "US pricing.",
    size=10,
)

# ---- 4. Design-partner terms ----
heading("4. Design-Partner (Pilot) Terms", 15)
para(
    "The current US prospect should be priced off the $2,499 Grow list, not around it — the "
    "discount must be framed so it protects the anchor rather than eroding it:"
)
table(
    ["Element", "Recommendation"],
    [
        ("Frame", "\u201cFounding design partner\u201d — a named program, not a discount."),
        ("Price", "$1,499/mo (~40% off the $2,499 Grow list) on a 12-month agreement; list price "
         "stated on the order form with the program credit shown against it. The all-modules-on "
         "decision puts payroll in scope, supporting the higher figure."),
        ("Exchange", "Weekly feedback sessions, US case study + logo rights, 2 reference calls, "
         "input on the US roadmap."),
        ("Expiry", "Program pricing holds 24 months, then steps to then-current list minus a "
         "loyalty tier — stated in the agreement so there is no renewal shock."),
        ("Scarcity", "Capped at 3-5 US design partners. Scarcity is what keeps a 40% concession "
         "from becoming the street price."),
    ],
)

# ---- 5. Open items ----
heading("5. Open Items", 15)
para(
    "1. Build the US pricing page: five columns (Start / Launch / Grow / Scale / Custom), stage "
    "+ size descriptions, and a \u201creplace your stack\u201d comparison strip.",
    size=10, after=2,
)
para(
    "2. Validate the stack-basket figures against the pilot client's actual current spend during "
    "discovery — their real invoice total is the best sales asset available.",
    size=10, after=2,
)
para(
    "3. Confirm the Start platform-fee mechanics (what the $99/mo covers, and the free-app usage "
    "limits that trigger a move to Launch).",
    size=10, after=2,
)
para(
    "4. Lock the exact band thresholds in billing (10 / 40 / 75 / 250 employees are the working "
    "defaults) and confirm the ~25-35% discount-to-stack band applied at each tier.",
    size=10,
)

# ---- Bottom line ----
heading("Bottom line", 15, GREEN)
para(
    "US list (flat per-org, whole platform included, 100% tier-matched to India, tuned so implied "
    "per-seat falls monotonically): Start — app free + $99/mo platform fee, self-serve; Launch "
    "$1,600 (10-40 emp); Grow $2,499 (40-75 emp); Scale $7,000 (75-250); Custom above 250; annual "
    "default. Each band is anchored to US market reality and priced ~25-35% below the stack it "
    "replaces. Priced by stage and size, never by module. The India book is untouched and never "
    "cross-referenced. The pilot client lands as a founding design partner at $1,499/mo (payroll "
    "in scope), with the discount explicitly purchased through feedback, references, and a case "
    "study. All figures are working numbers pending validation against real prospect spend.",
    italic=True, color=GREY, size=10,
)

doc.save(OUT)
print(f"Wrote {OUT}")
