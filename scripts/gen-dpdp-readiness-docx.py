#!/usr/bin/env python3
"""Generate the DPDP Readiness Checklist DOCX.

For the founder + privacy counsel. Splits readiness into (A) technical/code items
we can verify in the product, and (B) organisational/legal items that sit outside
the software and are required for an actual compliance posture. Each item is
marked DONE / FLAG-OFF / PARTIAL / NOT STARTED / COUNSEL. Deliberately honest:
the software provides DPDP-aligned tooling, not "compliance".
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/kathikiyer/Documents/NexusOps/docs-word/CoheronConnect_DPDP_Readiness_Checklist_2026-07-13.docx"

NAVY = RGBColor(0x1F, 0x33, 0x55)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6E, 0x00)
RED = RGBColor(0xC0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, italic=False, color=None, size=11, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def status_run(cell, value):
    p = cell.paragraphs[0]
    r = p.add_run(value)
    r.bold = True
    r.font.size = Pt(9.5)
    if value == "DONE":
        r.font.color.rgb = GREEN
    elif value in ("PARTIAL", "FLAG-OFF"):
        r.font.color.rgb = AMBER
    elif value == "COUNSEL":
        r.font.color.rgb = ACCENT
    else:  # NOT STARTED
        r.font.color.rgb = RED


def table(headers, rows, status_col):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i == status_col:
                status_run(cells[i], val)
            else:
                r = cells[i].paragraphs[0].add_run(val)
                r.font.size = Pt(9.5)
    return t


# ---- Title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("CoheronConnect")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
r = st.add_run("DPDP Act 2023 — Readiness Checklist")
r.font.size = Pt(14)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
r = meta.add_run(
    "Prepared 13 July 2026  •  For the founder and privacy counsel  •  "
    "Tracks both what the software does and the organisational/legal steps required. "
    "Technical items were verified against the source code."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()

# ---- Framing ----
heading("How to read this", 15)
para(
    "\"DPDP compliance\" is a property of how the organisation operates — not of the "
    "software alone. The product provides DPDP-aligned tooling (consent, data-subject "
    "requests, breach register, automated compliance sweeps). Turning that tooling into "
    "compliance requires enabling the remaining technical items AND completing the "
    "organisational/legal items in Section B, with counsel sign-off.",
)
para("Status legend: "
     "DONE = built & verified · "
     "PARTIAL = partly built · "
     "FLAG-OFF = built but disabled pending sign-off · "
     "NOT STARTED = not built · "
     "COUNSEL = legal/org action, outside code.",
     italic=True, color=GREY, size=9.5)

# ---- Section A: technical ----
heading("Section A — Technical / product items", 15)

heading("A1. Consent (DPDP §6)", 13, ACCENT, space_before=8, space_after=4)
table(
    ["Item", "Status", "Notes"],
    [
        ("Consent ledger: grant / withdraw / renew / expire, per (principal, purpose)", "DONE",
         "Append-only ledger; withdrawal as easy as grant."),
        ("Automatic expiry of lapsed consents", "DONE",
         "Scheduled sweep expires granted consents past their expiry."),
        ("Principal-facing consent capture UI / notice", "PARTIAL",
         "Backend exists; confirm front-end capture and §5 notice wording are wired."),
    ],
    status_col=1,
)

heading("A2. Data-subject rights (DPDP §11–14)", 13, ACCENT, space_before=8, space_after=4)
table(
    ["Item", "Status", "Notes"],
    [
        ("DSR intake with statutory response clock", "DONE",
         "Due date derived from receipt + response window."),
        ("DSR state machine + append-only event trail", "DONE",
         "Guarded transitions; full audit timeline."),
        ("Overdue-DSR alerting (automation)", "DONE",
         "Scheduled sweep raises internal alerts; deduped."),
        ("Right-to-erasure EXECUTION across all data stores", "FLAG-OFF",
         "Executor exists but disabled; erasure map covers 2 tables only. Needs counsel sign-off + full map."),
        ("Requestor identity verification", "NOT STARTED",
         "No verifiable identity check on DSR requestors yet."),
        ("Principal-facing DSR intake portal", "PARTIAL",
         "Confirm external submission channel is exposed and access-controlled."),
    ],
    status_col=1,
)

heading("A3. Breach handling (DPDP §8(6))", 13, ACCENT, space_before=8, space_after=4)
table(
    ["Item", "Status", "Notes"],
    [
        ("Breach register with notification clock", "DONE",
         "Notify-due derived from detection + jurisdiction window (72h default)."),
        ("Breach state machine + event trail", "DONE",
         "Board / principals notification stamps tracked."),
        ("Overdue-breach notice generation (automation)", "DONE",
         "Sweep generates board + principal notices; deduped."),
        ("ACTUAL delivery of notices (Board / principals)", "NOT STARTED",
         "Notices are logged only; no real email/SMS/Board channel delivery yet."),
    ],
    status_col=1,
)

heading("A4. Security controls supporting DPDP", 13, ACCENT, space_before=8, space_after=4)
table(
    ["Item", "Status", "Notes"],
    [
        ("Access control (RBAC) over personal data", "DONE", "Enforced per request."),
        ("Tamper-evident audit trail", "DONE", "SHA-256 hash chain + verifier."),
        ("Encrypted credentials / secrets", "DONE", "AES-256 for integration secrets."),
        ("Field-level encryption of personal data at rest", "NOT STARTED",
         "Personal-data columns stored in plaintext."),
        ("Database row-level security (defence in depth)", "NOT STARTED",
         "Tenant isolation is app-layer only."),
        ("Two-factor authentication (TOTP)", "PARTIAL",
         "Enrollment flag + step-up exist; TOTP not wired."),
        ("Data retention & deletion policy enforced in product", "NOT STARTED",
         "No automated retention/deletion schedule beyond consent expiry."),
    ],
    status_col=1,
)

# ---- Section B: organisational / legal ----
heading("Section B — Organisational / legal items (counsel-led)", 15)
para("These sit outside the software and are required for an actual compliance posture.",
     italic=True, color=GREY, size=9.5)
table(
    ["Item", "Status", "Notes"],
    [
        ("Appoint Data Protection Officer / point of contact (§13)", "COUNSEL",
         "Named, published contactable channel."),
        ("Published grievance-redressal mechanism (§13)", "COUNSEL",
         "Principal-facing route with response SLA."),
        ("Privacy notice & consent wording reviewed (§5–6)", "COUNSEL",
         "Plain-language, itemised purposes, withdrawal method."),
        ("Records of Processing Activities (RoPA) maintained", "COUNSEL",
         "Kept current; linked to consent/DSR records."),
        ("Data-processing agreements with all sub-processors (§8)", "COUNSEL",
         "Every third party touching personal data."),
        ("Retention & deletion policy (documented)", "COUNSEL",
         "Feeds the product retention item in A4."),
        ("Cross-border transfer position", "COUNSEL",
         "Where data is hosted/processed; permitted-country stance."),
        ("Sign-off on the erasure map (delete vs anonymise per table)", "COUNSEL",
         "Blocks enabling the flag-off erasure executor."),
        ("Alignment with notified DPDP Rules & timelines", "COUNSEL",
         "Configure statutory windows to the operative Rules once notified."),
        ("Breach notification runbook (who/when/how)", "COUNSEL",
         "Operational process behind the product's clock + notices."),
    ],
    status_col=1,
)

# ---- Go-live gate ----
heading("Minimum gate before any \"DPDP-ready\" claim", 15)
para("The following must be true before making even a soft public claim:", after=4)
def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
para("", after=2)
numbered("Erasure executor enabled with a counsel-approved, full erasure map (A2).")
numbered("Breach & DSR notices actually delivered, not just logged (A3, A2).")
numbered("DPO appointed and grievance channel published (B).")
numbered("Privacy notice / consent wording and RoPA signed off by counsel (B).")
numbered("Retention/deletion policy documented and, ideally, enforced in-product (A4/B).")

para("Until then, the honest public position is: \"CoheronConnect provides DPDP-aligned "
     "privacy tooling\" — not \"CoheronConnect is DPDP compliant.\"",
     italic=True, color=GREY, size=10)

doc.save(OUT)
print(f"Wrote {OUT}")
